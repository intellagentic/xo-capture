"""
XO Platform - Results Lambda
GET /results/:id — Returns analysis results for a client
POST /results/:id/brief — Downloads deployment brief as .docx
"""

import json
import os
import io
import re
import base64
import boto3
from datetime import datetime, timezone
from auth_helper import require_auth, get_db_connection, CORS_HEADERS, log_activity
try:
    from crypto_helper import unwrap_client_key, decrypt_s3_body
except ImportError:
    def unwrap_client_key(x): return None
    def decrypt_s3_body(k, b): return b if isinstance(b, str) else b.decode('utf-8', errors='replace') if b else ''

s3_client = boto3.client('s3')
BUCKET_NAME = os.environ.get('BUCKET_NAME', 'xo-client-data-mv')


def _get_enrichment_results(client_id, user):
    """Fetch latest enrichment results for a client. Returns (results_dict, error_response)."""
    conn = get_db_connection()
    cur = conn.cursor()

    is_admin = user.get('is_admin', False) or user.get('role') == 'admin'

    if is_admin:
        cur.execute("""
            SELECT e.status, e.results_s3_key, e.stage, c.encryption_key
            FROM enrichments e
            JOIN clients c ON e.client_id = c.id
            WHERE c.s3_folder = %s
            ORDER BY e.started_at DESC
            LIMIT 1
        """, (client_id,))
    else:
        cur.execute("""
            SELECT e.status, e.results_s3_key, e.stage, c.encryption_key
            FROM enrichments e
            JOIN clients c ON e.client_id = c.id
            WHERE c.s3_folder = %s AND c.user_id = %s
            ORDER BY e.started_at DESC
            LIMIT 1
        """, (client_id, user['user_id']))

    row = cur.fetchone()
    cur.close()
    conn.close()

    ck = unwrap_client_key(row[3]) if row and row[3] else None

    if row:
        enrichment_status, results_s3_key, enrichment_stage = row[0], row[1], row[2]

        if enrichment_status == 'processing':
            return None, {
                'statusCode': 200,
                'headers': CORS_HEADERS,
                'body': json.dumps({
                    'status': 'processing',
                    'stage': enrichment_stage or 'extracting',
                    'message': 'Analysis in progress'
                })
            }

        if enrichment_status == 'error':
            return None, {
                'statusCode': 200,
                'headers': CORS_HEADERS,
                'body': json.dumps({'status': 'error', 'message': 'Enrichment failed'})
            }

        s3_key = results_s3_key or f"{client_id}/results/analysis.json"
    else:
        s3_key = f"{client_id}/results/analysis.json"

    try:
        s3_response = s3_client.get_object(Bucket=BUCKET_NAME, Key=s3_key)
        raw = s3_response['Body'].read()
        decrypted = decrypt_s3_body(ck, raw)
        results = json.loads(decrypted)
        results['status'] = 'complete'
        return results, None
    except s3_client.exceptions.NoSuchKey:
        return None, {
            'statusCode': 200,
            'headers': CORS_HEADERS,
            'body': json.dumps({'status': 'processing', 'message': 'Analysis in progress'})
        }


def lambda_handler(event, context):
    if event.get('httpMethod') == 'OPTIONS':
        return {'statusCode': 200, 'headers': CORS_HEADERS, 'body': ''}

    user, err = require_auth(event)
    if err:
        log_activity(event, err)
        return err

    path = event.get('path', '')

    if '/brief' in path and event.get('httpMethod') == 'POST':
        response = _handle_brief_download(event, user)
    else:
        response = _handle_results(event, user)

    log_activity(event, response, user)
    return response


def _handle_results(event, user):
    try:
        path_params = event.get('pathParameters', {})
        client_id = path_params.get('id', '').strip()

        if not client_id:
            return {
                'statusCode': 400,
                'headers': CORS_HEADERS,
                'body': json.dumps({'error': 'client_id is required'})
            }

        results, error_resp = _get_enrichment_results(client_id, user)
        if error_resp:
            return error_resp

        return {
            'statusCode': 200,
            'headers': CORS_HEADERS,
            'body': json.dumps(results)
        }

    except Exception as e:
        print(f"Error retrieving results: {str(e)}")
        return {
            'statusCode': 500,
            'headers': CORS_HEADERS,
            'body': json.dumps({'error': 'Internal server error', 'message': str(e)})
        }


# ── Deployment Brief Assembly ──

def _assemble_brief(results, client_name='', industry='', description='', contact_name=''):
    """Assemble deployment brief from existing analysis results — no Bedrock call."""
    problems = results.get('problems', results.get('problems_identified', []))
    primary = problems[0] if problems else {}
    plan = results.get('plan', results.get('action_plan', {}))
    plan_phases = []
    if isinstance(plan, list):
        plan_phases = plan
    elif isinstance(plan, dict):
        plan_phases = [{'phase': k, 'actions': v} for k, v in plan.items()]

    company = client_name or results.get('company_name', 'the client')
    ind = industry or 'this domain'

    return {
        'cover': {
            'client_name': company,
            'client_descriptor': description or ind,
            'headline': f"XO Deployment: {primary.get('title', 'Operational Transformation')}" if primary else f"XO Deployment for {company}",
            'value_proposition': (results.get('bottom_line', '') or '').split('.')[0] + '.' if results.get('bottom_line') else '',
            'client_contact': contact_name,
            'meeting_date': 'TBD',
        },
        'executive_summary': results.get('summary', results.get('executive_summary', '')),
        'key_metrics': [
            {'value': str(len(problems)), 'label': 'Critical Issues Identified',
             'sublabel': f"{sum(1 for p in problems if p.get('severity') == 'high')} high severity"},
        ] + [{'value': p.get('severity', 'N/A').upper(), 'label': p.get('title', '')[:40], 'sublabel': f"{p.get('severity', '')} priority"} for p in problems[:2]],
        'sections': [
            {'number': '01', 'title': f'CLIENT PROFILE: {company}',
             'content': (results.get('summary', '') or '') + (f"\n\n**Industry:** {ind}\n\n{description}" if description else ''),
             'callout': {'label': f'THE {ind.upper()} CONTEXT', 'content': primary.get('evidence', '')}},
            {'number': '02', 'title': 'THE OPERATIONAL CRISIS',
             'content': '\n\n---\n\n'.join(
                 f"**{p.get('title', '')}** ({p.get('severity', '')} severity)\n{p.get('evidence', '')}\n\n**Recommendation:** {p.get('recommendation', '')}"
                 for p in problems)},
            {'number': '03', 'title': 'WHY STANDARD AI CANNOT BE USED HERE',
             'content': f"Generic AI tools like ChatGPT or off-the-shelf automation platforms cannot safely operate in {ind} because they lack domain-specific guardrails. In {company}'s environment, a single error in {primary.get('title', 'operational processes').lower() if primary else 'operational processes'} could result in {primary.get('evidence', 'significant compliance and operational failures').split('.')[0].lower() if primary else 'significant compliance and operational failures'}.\n\nStandard AI has no concept of {ind} compliance hierarchies, cannot cross-reference domain-specific standards and regulations, and provides no audit trail for regulatory accountability. The XO platform's Constitutional Safety layer ensures that every AI-generated output is bounded by domain rules that the operator defines and controls."},
            {'number': '04', 'title': 'THE XO DEPLOYMENT: ARCHITECTURE & OODA WORKFLOW',
             'content': (f"```\n{results.get('architecture_diagram', '')}\n```\n\n" if results.get('architecture_diagram') else '') +
                f"The XO deployment for {company} operates on a continuous **Observe-Orient-Decide-Act** loop, processing {ind} data through domain-specific rules before any output reaches the operator.\n\n"
                f"**Observe:** XO ingests documents, data feeds, and operational inputs from {company}'s systems.\n"
                f"**Orient:** The DX Cartridge contextualises each input against {ind} rules, standards, and historical patterns.\n"
                f"**Decide:** XO generates recommendations bounded by Constitutional Safety rules — flagging items that require human judgment.\n"
                f"**Act:** Approved outputs are delivered through Streamline workflows, with full audit logging."},
            {'number': '05', 'title': 'CONSTITUTIONAL SAFETY',
             'content': f"XO enforces a Constitutional Layer — a set of immutable domain rules that the AI cannot override. For {company}, this means:\n\n"
                f"- **Compliance Validation:** Every output is validated against {ind} standards and regulations before delivery\n"
                f"- **Human Authority:** The operator retains final authority on all decisions flagged as requiring human judgment\n"
                f"- **Audit Trail:** All AI actions are logged with full provenance for regulatory audit\n"
                f"- **Domain Boundaries:** Boundaries are encoded as rules, not suggestions — the system cannot generate outputs that violate them"},
            {'number': '06', 'title': 'PROOF OF CONCEPT & NEXT STEPS',
             'content': '\n\n'.join(f"**{p.get('phase', '')}**\n" + '\n'.join(f"{i+1}. {re.sub(r'^\\d+\\.\\s*', '', a)}" for i, a in enumerate(p.get('actions', []) if isinstance(p.get('actions'), list) else []))
                                    for p in plan_phases)},
        ],
        'ooda_phases': [
            {'name': 'OBSERVE', 'tagline': f"Ingests {company}'s operational data", 'bullets': ['Document upload and extraction', 'Data feed integration', 'Historical pattern capture']},
            {'name': 'ORIENT', 'tagline': f'Contextualises against {ind} rules', 'bullets': ['Domain rule matching', 'Compliance cross-reference', 'Risk classification']},
            {'name': 'DECIDE', 'tagline': 'Generates bounded recommendations', 'bullets': ['AI analysis within safety constraints', 'Human-judgment flagging', 'Confidence scoring']},
            {'name': 'ACT', 'tagline': 'Delivers through Streamline workflows', 'bullets': ['Automated report generation', 'Notification and escalation', 'Full audit logging']},
        ],
        'poc_timeline': [
            {'step': '1', 'timeline': 'Week 1', 'action': re.sub(r'^\d+\.\s*', '', (plan_phases[0].get('actions', ['Configure DX Cartridge'])[0] if plan_phases else 'Configure DX Cartridge with domain rules'))},
            {'step': '2', 'timeline': 'Week 1-2', 'action': re.sub(r'^\d+\.\s*', '', (plan_phases[0].get('actions', ['', 'Ingest sample data'])[1] if plan_phases and len(plan_phases[0].get('actions', [])) > 1 else 'Ingest sample data and validate extraction'))},
            {'step': '3', 'timeline': 'Week 2', 'action': re.sub(r'^\d+\.\s*', '', (plan_phases[1].get('actions', ['Run analysis'])[0] if len(plan_phases) > 1 else 'Run analysis against live data'))},
            {'step': '4', 'timeline': 'Week 3', 'action': re.sub(r'^\d+\.\s*', '', (plan_phases[2].get('actions', ['Deploy decision'])[0] if len(plan_phases) > 2 else 'Review results and make deploy/iterate decision'))},
        ],
        'success_metric': f"The pilot is successful when {primary.get('title', 'the primary operational bottleneck').lower()} is resolved without manual intervention in the current workflow." if primary else 'The pilot is successful when the primary operational bottleneck is resolved through automated XO processing.',
    }


def _handle_brief_download(event, user):
    try:
        path_params = event.get('pathParameters', {})
        client_id = path_params.get('id', '').strip()

        if not client_id:
            return {
                'statusCode': 400,
                'headers': CORS_HEADERS,
                'body': json.dumps({'error': 'client_id is required'})
            }

        results, error_resp = _get_enrichment_results(client_id, user)
        if error_resp:
            return error_resp

        # Fetch client record for industry/description
        conn = get_db_connection()
        cur = conn.cursor()
        cur.execute("SELECT company_name, industry, description, contact_name FROM clients WHERE s3_folder = %s", (client_id,))
        client_row = cur.fetchone()
        cur.close()
        conn.close()

        client_name = client_row[0] if client_row else ''
        industry = client_row[1] if client_row else ''
        description = client_row[2] if client_row else ''
        contact_name = client_row[3] if client_row else ''

        brief = _assemble_brief(results, client_name, industry, description, contact_name)
        if not brief:
            return {
                'statusCode': 404,
                'headers': CORS_HEADERS,
                'body': json.dumps({'error': 'No deployment brief available for this analysis'})
            }

        company_name = results.get('company_name', brief.get('cover', {}).get('client_name', 'Client'))

        # Generate .docx
        docx_bytes = _generate_brief_docx(brief, company_name)

        filename = f"{company_name.replace(' ', '_')}_XO_Deployment_Brief.docx"

        return {
            'statusCode': 200,
            'headers': CORS_HEADERS,
            'body': json.dumps({
                'filename': filename,
                'content_base64': base64.b64encode(docx_bytes).decode('utf-8'),
                'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            })
        }

    except Exception as e:
        print(f"Error generating brief: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            'statusCode': 500,
            'headers': CORS_HEADERS,
            'body': json.dumps({'error': 'Failed to generate brief', 'message': str(e)})
        }


def _generate_brief_docx(brief, company_name):
    """Generate a .docx deployment brief from structured JSON."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    TEAL = RGBColor(0x0F, 0x96, 0x9C)
    DARK = RGBColor(0x1a, 0x1a, 0x2e)
    MUTED = RGBColor(0x66, 0x66, 0x66)

    # Cover
    cover = brief.get('cover', {})
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(6)
    run = p.add_run('XO DEPLOYMENT BRIEF')
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(4)
    run = p.add_run(cover.get('headline', f'XO Deployment for {company_name}'))
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(4)
    run = p.add_run(f"{cover.get('client_name', company_name)} — {cover.get('client_descriptor', '')}")
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    if cover.get('value_proposition'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(12)
        run = p.add_run(cover['value_proposition'])
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = TEAL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(24)
    run = p.add_run(f"Prepared by Intellagentic Limited | {datetime.now(timezone.utc).strftime('%B %d, %Y')}")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    # Executive Summary
    if brief.get('executive_summary'):
        p = doc.add_paragraph()
        run = p.add_run('EXECUTIVE SUMMARY')
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = TEAL
        p.space_after = Pt(8)

        p = doc.add_paragraph(brief['executive_summary'])
        p.style.font.size = Pt(11)
        p.paragraph_format.line_spacing = Pt(16)
        p.space_after = Pt(16)

    # Key Metrics
    metrics = brief.get('key_metrics', [])
    if metrics:
        table = doc.add_table(rows=1, cols=len(metrics))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, m in enumerate(metrics):
            cell = table.cell(0, i)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(m.get('value', ''))
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = TEAL
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(m.get('label', ''))
            run.font.size = Pt(9)
            run.font.bold = True
            if m.get('sublabel'):
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(m['sublabel'])
                run.font.size = Pt(8)
                run.font.color.rgb = MUTED
        doc.add_paragraph()  # spacer

    # Numbered Sections
    for sec in brief.get('sections', []):
        p = doc.add_paragraph()
        p.space_before = Pt(16)
        p.space_after = Pt(8)
        run = p.add_run(f"{sec.get('number', '')}  ")
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = TEAL
        run = p.add_run(sec.get('title', ''))
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = DARK

        content = sec.get('content', '')
        # Strip backslash escapes
        content = content.replace('\\*', '*').replace('\\-', '-').replace('\\|', '|')
        # Split on fenced code blocks
        segments = re.split(r'(```[\s\S]*?```)', content)
        for segment in segments:
            if segment.startswith('```') and segment.endswith('```'):
                # Code block — render as monospace preformatted
                code = segment.strip('`').strip()
                if code.startswith('\n'): code = code[1:]
                # Remove language hint (e.g. "python\n")
                if '\n' in code and not code.split('\n')[0].strip().startswith(('+', '|', '-', ' ')):
                    first_line = code.split('\n')[0].strip()
                    if len(first_line) < 20 and ' ' not in first_line:
                        code = code[len(first_line):].lstrip('\n')
                for line in code.split('\n'):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing = Pt(12)
                    run = p.add_run(line)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(7)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            else:
                # Regular text
                for para_text in segment.split('\n\n'):
                    stripped = para_text.strip()
                    if not stripped:
                        continue
                    # Strip markdown bold for plain text
                    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
                    p = doc.add_paragraph(clean)
                    p.paragraph_format.line_spacing = Pt(16)

        callout = sec.get('callout')
        if callout:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f"{callout.get('label', '')}: ")
            run.font.bold = True
            run.font.color.rgb = TEAL
            run.font.size = Pt(10)
            run = p.add_run(callout.get('content', ''))
            run.font.size = Pt(10)
            run.font.italic = True

    # OODA Phases
    ooda = brief.get('ooda_phases', [])
    if ooda:
        p = doc.add_paragraph()
        p.space_before = Pt(20)
        run = p.add_run('OODA WORKFLOW')
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = TEAL

        for phase in ooda:
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            run = p.add_run(f"{phase.get('name', '')} — ")
            run.font.bold = True
            run.font.color.rgb = TEAL
            run.font.size = Pt(11)
            run = p.add_run(phase.get('tagline', ''))
            run.font.size = Pt(11)
            run.font.italic = True

            for bullet in phase.get('bullets', []):
                p = doc.add_paragraph(bullet, style='List Bullet')
                p.style.font.size = Pt(10)

    # POC Timeline
    poc = brief.get('poc_timeline', [])
    if poc:
        p = doc.add_paragraph()
        p.space_before = Pt(20)
        p.space_after = Pt(8)
        run = p.add_run('PROOF OF CONCEPT TIMELINE')
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = TEAL

        table = doc.add_table(rows=len(poc) + 1, cols=3)
        table.style = 'Table Grid'
        # Header
        for i, label in enumerate(['Step', 'Timeline', 'Action']):
            cell = table.cell(0, i)
            cell.paragraphs[0].add_run(label).bold = True
        # Rows
        for i, row in enumerate(poc):
            table.cell(i + 1, 0).text = str(row.get('step', ''))
            table.cell(i + 1, 1).text = row.get('timeline', '')
            table.cell(i + 1, 2).text = row.get('action', '')

    # Success Metric
    if brief.get('success_metric'):
        p = doc.add_paragraph()
        p.space_before = Pt(20)
        run = p.add_run('SUCCESS METRIC: ')
        run.font.bold = True
        run.font.color.rgb = TEAL
        run.font.size = Pt(11)
        run = p.add_run(brief['success_metric'])
        run.font.size = Pt(11)

    # Footer
    p = doc.add_paragraph()
    p.space_before = Pt(30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Intellagentic Limited | Company No. 16761110 | intellagentic.io')
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    # Write to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
