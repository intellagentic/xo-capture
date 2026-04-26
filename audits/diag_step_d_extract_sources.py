"""Step D part 1: pull each source file from S3, decrypt with the per-client
key, and extract text using the same code paths xo-enrich uses. Dump the
first 5500 chars of each file (Claude saw the first 5000 — a buffer of 500
shows what was just past the cap)."""
import os
import sys
import io
import csv
import psycopg2
import boto3

sys.path.insert(0, '/Users/ken_macair_2025/xo-capture/backend/lambdas/enrich')
if 'AES_MASTER_KEY' not in os.environ:
    raise SystemExit("AES_MASTER_KEY must be set in env before running")

from crypto_helper import unwrap_client_key, maybe_decrypt_s3_bytes  # noqa

DB_URL = os.environ['DATABASE_URL']
BUCKET = 'xo-client-data-mv'

session = boto3.Session(profile_name='intellagentic', region_name='eu-west-2')
s3 = session.client('s3', region_name='us-west-1')

CLIENTS = [
    ('FC Dynamics',          '51f49469-6328-4afe-a492-7c2f36274907', 'client_1772616693_8b881fe7'),
    ('MFP Trading Limited',  '58420e26-fd85-4da7-a638-e8729b55725f', 'client_1776011770_6aff114c'),
]


def extract_csv(b):
    content = b.decode('utf-8', errors='replace')
    reader = csv.reader(io.StringIO(content))
    rows = list(reader)
    text = f"CSV Data:\nTotal rows: {len(rows)}\n\n"
    if rows:
        text += "Header: " + ", ".join(rows[0]) + "\n\n"
        text += "Sample data (first 10 rows):\n"
        for i, row in enumerate(rows[1:11]):
            text += f"Row {i+1}: " + ", ".join(row) + "\n"
    return text


def extract_excel(b):
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(b), data_only=True)
    text = "Excel Data:\n\n"
    for s_ in wb.sheetnames:
        sh = wb[s_]
        text += f"Sheet: {s_}\nRows: {sh.max_row}, Columns: {sh.max_column}\n\n"
        text += "Sample data (first 10 rows):\n"
        for i, row in enumerate(sh.iter_rows(max_row=10, values_only=True)):
            text += f"Row {i+1}: " + ", ".join(str(c) if c else "" for c in row) + "\n"
        text += "\n"
    return text


def extract_pdf(b):
    from pypdf import PdfReader
    pdf = PdfReader(io.BytesIO(b))
    text = f"PDF Document ({len(pdf.pages)} pages):\n\n"
    for i, page in enumerate(pdf.pages[:10]):
        page_text = page.extract_text()
        if page_text:
            text += f"--- Page {i+1} ---\n{page_text}\n\n"
    return text


def extract_docx(b):
    from docx import Document
    doc = Document(io.BytesIO(b))
    text = f"Word Document ({len(doc.paragraphs)} paragraphs):\n\n"
    for p in doc.paragraphs:
        if p.text.strip():
            text += p.text + "\n"
    if doc.tables:
        text += f"\n--- Tables ({len(doc.tables)}) ---\n"
        for i, t in enumerate(doc.tables):
            text += f"\nTable {i+1}:\n"
            for r in t.rows:
                text += " | ".join(c.text.strip() for c in r.cells) + "\n"
    return text


def extract_text(filename, b):
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    if ext == 'csv':   return extract_csv(b)
    if ext == 'txt':   return b.decode('utf-8', errors='replace')
    if ext in ('xlsx', 'xls'): return extract_excel(b)
    if ext == 'pdf':   return extract_pdf(b)
    if ext in ('docx', 'doc'): return extract_docx(b)
    return f"<unsupported ext: {ext}>"


conn = psycopg2.connect(DB_URL)

for client_name, db_id, s3_folder in CLIENTS:
    print(f"\n{'='*80}\n{client_name}\n{'='*80}")

    cur = conn.cursor()
    cur.execute("SELECT encryption_key FROM clients WHERE id = %s", (db_id,))
    enc_key_b64 = cur.fetchone()[0]
    cur.execute("""
        SELECT s3_key, filename, uploaded_at
        FROM uploads
        WHERE client_id = %s AND status = 'active'
        ORDER BY s3_key
    """, (db_id,))
    rows = cur.fetchall()
    cur.close()

    client_key = unwrap_client_key(enc_key_b64)

    out_dir = f"/Users/ken_macair_2025/xo-capture/audits/source-text/{s3_folder}"
    os.makedirs(out_dir, exist_ok=True)

    for s3_key, filename, uploaded_at in rows:
        try:
            obj = s3.get_object(Bucket=BUCKET, Key=s3_key)
            raw = obj['Body'].read()
            decrypted = maybe_decrypt_s3_bytes(client_key, raw)
            text = extract_text(filename, decrypted)
            total = len(text)
            seen_by_claude = text[:5000]
            buffer_after = text[5000:5500]
            beyond_cap = text[5500:]

            safe_name = filename.replace('/', '_')[:200]
            out_path = f"{out_dir}/{safe_name}.txt"
            with open(out_path, 'w') as f:
                f.write(f"FILENAME: {filename}\n")
                f.write(f"UPLOADED: {uploaded_at}\n")
                f.write(f"TOTAL_EXTRACTED_CHARS: {total}\n")
                f.write(f"CLAUDE_SAW: first 5000 chars (Lambda truncates with text[:5000])\n")
                f.write(f"BEYOND_CAP_CHARS: {max(0, total - 5000)}\n\n")
                f.write("--- FIRST 5000 CHARS (CLAUDE SAW THIS) ---\n")
                f.write(seen_by_claude)
                f.write("\n\n--- CHARS 5000-5500 (BUFFER, JUST PAST CAP) ---\n")
                f.write(buffer_after)
                if beyond_cap:
                    f.write("\n\n--- CHARS 5500+ (CLAUDE DID NOT SEE) ---\n")
                    # only first 4000 of the rest, to keep file size sensible
                    f.write(beyond_cap[:4000])
                    if len(beyond_cap) > 4000:
                        f.write(f"\n\n[... {len(beyond_cap) - 4000} more chars truncated ...]")
            print(f"  wrote {filename!r:80}  total={total:>7}  seen=5000  beyond={max(0, total-5000)}")
        except Exception as e:
            print(f"  FAILED {filename!r}: {e}")

conn.close()
print("\nDone.")
